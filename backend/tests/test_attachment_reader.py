from app.assistant.attachment_reader import extract_text


def test_extract_md_text():
    out = extract_text("# 标题\n正文内容".encode("utf-8"), "说明.md", 20000)
    assert "正文内容" in out["text"]
    assert out["truncated"] is False
    assert out["file_name"] == "说明.md"


def test_extract_csv_text():
    out = extract_text("a,b\n1,2".encode("utf-8"), "data.CSV", 20000)
    assert "1,2" in out["text"]


def test_unsupported_format_returns_error():
    out = extract_text(b"\x89PNG....", "图片.png", 20000)
    assert "error" in out
    assert out["file_name"] == "图片.png"


def test_truncation():
    out = extract_text(("x" * 100).encode(), "big.txt", 10)
    assert out["truncated"] is True
    assert len(out["text"]) == 10
    assert out["chars"] == 100


import io


def _make_docx_bytes():
    import docx
    d = docx.Document()
    d.add_paragraph("会议纪要正文")
    t = d.add_table(rows=1, cols=2)
    t.rows[0].cells[0].text = "零件号"
    t.rows[0].cells[1].text = "P-100"
    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def _make_xlsx_bytes():
    import openpyxl
    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "清单"
    ws.append(["编码", "数量"])
    ws.append(["P-100", 3])
    buf = io.BytesIO()
    wb.save(buf)
    return buf.getvalue()


def test_extract_docx_paragraph_and_table():
    out = extract_text(_make_docx_bytes(), "纪要.docx", 20000)
    assert "会议纪要正文" in out["text"]
    assert "P-100" in out["text"]


def test_extract_xlsx_cells():
    out = extract_text(_make_xlsx_bytes(), "清单.xlsx", 20000)
    assert "清单" in out["text"]
    assert "P-100" in out["text"]


def test_corrupted_docx_returns_error():
    out = extract_text(b"not a real docx", "bad.docx", 20000)
    assert "error" in out
