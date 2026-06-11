"""附件正文提取：按扩展名分派（pdf/docx/xlsx/文本），供 AI 分析。

依赖（pypdf/python-docx/openpyxl）在分支内按需 import，避免导入期硬依赖。
"""
import io
import os

TEXT_EXTS = {".md", ".txt", ".csv", ".json"}


def _ext(file_name: str) -> str:
    return os.path.splitext(file_name or "")[1].lower()


def _extract_pdf(data: bytes) -> str:
    from pypdf import PdfReader
    reader = PdfReader(io.BytesIO(data))
    return "\n".join((page.extract_text() or "") for page in reader.pages)


def _extract_docx(data: bytes) -> str:
    import docx
    d = docx.Document(io.BytesIO(data))
    parts = [p.text for p in d.paragraphs]
    for table in d.tables:
        for row in table.rows:
            parts.append("\t".join(c.text for c in row.cells))
    return "\n".join(parts)


def _extract_xlsx(data: bytes) -> str:
    import openpyxl
    wb = openpyxl.load_workbook(io.BytesIO(data), read_only=True, data_only=True)
    parts = []
    for ws in wb.worksheets:
        parts.append(f"[sheet] {ws.title}")
        for row in ws.iter_rows(values_only=True):
            parts.append("\t".join("" if v is None else str(v) for v in row))
    return "\n".join(parts)


_EXTRACTORS = {".pdf": _extract_pdf, ".docx": _extract_docx, ".xlsx": _extract_xlsx}


def extract_text(data: bytes, file_name: str, max_chars: int) -> dict:
    """提取附件正文。返回 {file_name, text, truncated, chars} 或 {file_name, error}。"""
    ext = _ext(file_name)
    try:
        if ext in TEXT_EXTS:
            text = data.decode("utf-8", errors="replace")
        elif ext in _EXTRACTORS:
            text = _EXTRACTORS[ext](data)
        else:
            return {"file_name": file_name,
                    "error": f"该格式（{ext or '未知'}）暂不支持提取正文，"
                             "支持 pdf/docx/xlsx/md/txt/csv/json"}
    except Exception as exc:
        return {"file_name": file_name, "error": f"提取失败: {exc}"}
    chars = len(text)
    truncated = chars > max_chars
    return {"file_name": file_name, "text": text[:max_chars],
            "truncated": truncated, "chars": chars}
